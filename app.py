import streamlit as st
import google.generativeai as genai
import pandas as pd
from docx import Document
from io import BytesIO

# ---- Configure Gemini API ----
GENAI_API_KEY = st.secrets["GEMINI_API_KEY"]
genai.configure(api_key=GENAI_API_KEY)

# ---- Initialize Session State for History ----
if "prediction_history" not in st.session_state:
    st.session_state.prediction_history = []

# ---- Function to get FIR sections and descriptions ----
def get_fir_sections(case_description):
    model = genai.GenerativeModel("gemini-2.0-flash")
    
    prompt = (
        "You are an expert in the Bhartiya Nyaya Sanhita (BNS), 2023. Given a case description, "
        "identify the most relevant legal sections from BNS 2023 and provide their descriptions. "
        "Return data in the following format:\n\n"
        "**Format:**\n"
        "Section 101: [Brief description of section 101]\n"
        "Section 103: [Brief description of section 103]\n"
        "Section 112: [Brief description of section 112]\n\n"
        "**Example:**\n"
        "Case: A person was attacked with a deadly weapon, causing severe injury.\n"
        "Output:\n"
        "Section 101: Attempt to cause grievous hurt with a dangerous weapon.\n"
        "Section 103: Voluntarily causing hurt with intent to disable.\n"
        "Section 112: Punishment for causing grievous hurt.\n\n"
        "Now, analyze this case and return only relevant BNS sections along with their descriptions:\n\n"
        f"Case: {case_description}"
    )

    response = model.generate_content(prompt)
    return response.text.strip() if response else "Error fetching response"

# ---- Function to Create Word File ----
def create_word_file(case_description, prediction_result):
    doc = Document()
    doc.add_heading("FIR Section Prediction Report", level=1)

    # Add case description
    doc.add_heading("Case Description:", level=2)
    doc.add_paragraph(case_description)

    # Add Predicted Sections
    doc.add_heading("Predicted Sections & Descriptions:", level=2)
    
    lines = prediction_result.split("\n")
    for line in lines:
        if line.startswith("Section"):
            parts = line.split(": ")
            if len(parts) == 2:
                section, description = parts
                para = doc.add_paragraph()
                para.add_run(section + ": ").bold = True  # Highlight section number
                para.add_run(description)
        else:
            doc.add_paragraph(line)

    # Save to BytesIO for download
    word_buffer = BytesIO()
    doc.save(word_buffer)
    word_buffer.seek(0)
    return word_buffer

# ---- Streamlit UI ----
st.set_page_config(page_title="FIR Section Predictor", page_icon="‚öñÔ∏è", layout="wide")

# ---- Sidebar ----
with st.sidebar:
    st.title("‚öñÔ∏è FIR Section Predictor")
    st.markdown("### A legal AI tool to predict relevant FIR sections.")
    st.markdown("---")
    st.info("üöÄ **Usage:** Enter a case description and click **Predict Sections** button.")

    # Clear History Button
    if st.button("üóëÔ∏è Clear History"):
        st.session_state.prediction_history = []
        st.success("History cleared successfully!")

st.markdown("## üîç Enter Case Description")
st.write("Provide a brief case description, and the system will predict the relevant **Bhartiya Nyaya Sanhita (BNS) 2023** sections.")

# Text Area for Description
case_description = st.text_area("‚úçÔ∏è Case Description:", height=150, placeholder="Type case details here...")

# Predict Button
predict_btn = st.button("‚ö° Predict Sections", use_container_width=True)

if predict_btn:
    if case_description.strip():
        with st.spinner("üîé Analyzing case details..."):
            prediction_result = get_fir_sections(case_description)

        st.success("‚úÖ Prediction Complete!")
        st.subheader("üìå Predicted Sections & Descriptions:")

        # Display Result with Highlighted Section Numbers
        for line in prediction_result.split("\n"):
            if line.startswith("Section"):
                parts = line.split(": ")
                if len(parts) == 2:
                    section, description = parts
                    st.markdown(f"**üìñ {section}:** {description}")
            else:
                st.write(line)

        # Save prediction in history
        st.session_state.prediction_history.append({
            "case": case_description,
            "prediction": prediction_result
        })

        # Generate Word File
        word_file = create_word_file(case_description, prediction_result)

        # Download Button for Word File
        st.download_button(
            label="üì• Download Predictions as Word File",
            data=word_file,
            file_name="FIR_Predictions.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    else:
        st.warning("‚ö†Ô∏è Please enter a case description before predicting.")

# ---- History Section ----
if st.session_state.prediction_history:
    with st.expander("üìú View Previous Predictions"):
        for i, entry in enumerate(reversed(st.session_state.prediction_history), start=1):
            st.markdown(f"### üìù Case {i}")
            st.write(f"**üîπ Case Description:** {entry['case']}")
            st.markdown("**üîπ Predicted Sections & Descriptions:**")
            for line in entry["prediction"].split("\n"):
                if line.startswith("Section"):
                    parts = line.split(": ")
                    if len(parts) == 2:
                        section, description = parts
                        st.markdown(f"**üìñ {section}:** {description}")
                else:
                    st.write(line)
            st.markdown("---")

# Footer
st.markdown("---")
st.markdown("üìú **Note:** This tool provides AI-generated predictions and should be verified by legal professionals.")
